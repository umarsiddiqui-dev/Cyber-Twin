from docx import Document
import json

def extract_text(filename):
    doc = Document(filename)
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text.append(" | ".join(row_text))
            
    with open(f"{filename}.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(text))

extract_text('CyberTwin-FYDP Proposal-FINALIZED.docx')
extract_text('Starting pages of Reports Batch 2022F (1).docx')
extract_text('FYDP Report Format Updated for 2022F (1).docx')
